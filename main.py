import streamlit as st
import PyPDF2
from docx import Document
from PIL import Image
import pytesseract
import easyocr
import os
from datetime import datetime
from extraction import ThreeLayerExtractor
import json

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

st.set_page_config(
    page_title="FinanceInsight - Financial Document Intelligence",
    page_icon="💰",
    layout="wide"
)


if "extractor" not in st.session_state:
    with st.spinner("Loading FinBERT + Gemini..."):
        st.session_state.extractor = ThreeLayerExtractor()

if "extraction_result" not in st.session_state:
    st.session_state.extraction_result = None

if "last_extracted_file" not in st.session_state:
    st.session_state.last_extracted_file = None

if "extraction_in_progress" not in st.session_state:
    st.session_state.extraction_in_progress = False

if "ocr_reader" not in st.session_state:
    with st.spinner("Loading OCR engine..."):
        try:
            st.session_state.ocr_reader = easyocr.Reader(['en'], gpu=True)
            print("✓ EasyOCR using GPU")
        except:
            st.session_state.ocr_reader = easyocr.Reader(['en'], gpu=False)
            print("EasyOCR using CPU (GPU not available)")

MAX_CHARACTERS = 5000


def extract_text_from_image_pytesseract(image_file):
    try:
        img = Image.open(image_file)
        if img.mode in ('RGBA', 'LA', 'P'):
            img = img.convert('RGB')
        text = pytesseract.image_to_string(img, lang='eng')
        return text
    except Exception as e:
        raise Exception(f"Tesseract OCR error: {str(e)}")

def extract_text_from_image_easyocr(image_file):
    try:
        img = Image.open(image_file)
        
        if img.mode == 'RGBA':
            img = img.convert('RGB')
        elif img.mode == 'LA':
            img = img.convert('L')
        elif img.mode == 'P':
            img = img.convert('RGB')
        
        temp_path = "temp_image.jpg"
        img.save(temp_path, "JPEG", quality=95)  
        
        result = st.session_state.ocr_reader.readtext(temp_path, detail=0)
        text = "\n".join(result)
        
        if os.path.exists(temp_path):
            os.remove(temp_path)
        
        return text
    except Exception as e:
        raise Exception(f"EasyOCR error: {str(e)}")

def extract_text_from_file(file):
    try:
        file_type = file.type
        file_name = file.name
        
        if file_type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text: 
                    text += page_text + "\n"
            return text
        
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():  
                    text += para.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            return text
        
        elif file_type == "text/plain":
            text = file.read().decode("utf-8")
            return text
        
        elif file_type in ["image/png", "image/jpeg", "image/jpg", "image/bmp", "image/tiff"]:
            print(f"Detected image file: {file_name}")
            
            try:
                text = extract_text_from_image_easyocr(file)
                print(f"✓ EasyOCR extracted {len(text)} characters")
                return text
            except Exception as easyocr_error:
                print(f"EasyOCR failed: {easyocr_error}")
                try:
                    text = extract_text_from_image_pytesseract(file)
                    print(f"✓ Tesseract extracted {len(text)} characters")
                    return text
                except Exception as tesseract_error:
                    raise Exception(
                        f"Both OCR engines failed:\n"
                        f"EasyOCR: {easyocr_error}\n"
                        f"Tesseract: {tesseract_error}"
                    )
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    except Exception as e:
        raise Exception(f"Error extracting text: {str(e)}")



st.title("💰 FinanceInsight")
st.subheader("Smart Financial Document Intelligence")


with st.sidebar:
    st.header("📁 Document Upload")
    
    uploaded_file = st.file_uploader(
        "Upload Financial Document",
        type=["pdf", "docx", "doc", "txt", "png", "jpg", "jpeg", "bmp", "tiff"],
        help="PDF, Word, Text, or Image file with financial information"
    )
    
    if uploaded_file:
        st.write(f"**File:** {uploaded_file.name}")
        st.write(f"**Type:** {uploaded_file.type}")
        st.write(f"**Size:** {uploaded_file.size / 1024:.1f} KB")
        
        if uploaded_file.type.startswith("image/"):
            st.image(uploaded_file, caption="Uploaded Image", use_container_width=True)

col1, col2 = st.columns([3, 2], gap="medium")


with col1:
    st.header("Entity Extraction")
    
    if uploaded_file:
        try:
            if uploaded_file.type.startswith("image/"):
                st.info("Image file detected - Using OCR for text extraction")
            
            with st.spinner("Extracting text..."):
                document_text = extract_text_from_file(uploaded_file)
            
            original_length = len(document_text)
            document_text = document_text[:MAX_CHARACTERS]
            
            if original_length > MAX_CHARACTERS:
                st.warning(f"Truncated to {MAX_CHARACTERS} characters")
            
            st.success(f"✓ Extracted {len(document_text)} characters")
            
            with st.expander("Document Preview"):
                st.text_area(
                    "Extracted Text:",
                    value=document_text[:500] + ("..." if len(document_text) > 500 else ""),
                    height=150,
                    disabled=True
                )
            
            if st.button("Extract Entities", use_container_width=True, key="extract"):
                
                if st.session_state.extraction_in_progress:
                    st.warning("Extraction already in progress. Please wait...")
                
                elif st.session_state.last_extracted_file == uploaded_file.name:
                    st.info("This document was already extracted. Showing cached results.")
                
                else:
                    st.session_state.extraction_in_progress = True
                    
                    with st.spinner("Extracting..."):
                        try:
                            extraction_result = st.session_state.extractor.extract_comprehensive(
                                document_text
                            )
                            
                            st.session_state.extraction_result = extraction_result
                            st.session_state.last_extracted_file = uploaded_file.name
                            
                            st.success("Extraction Complete!")
                        
                        except Exception as e:
                            st.error(f" Error: {str(e)}")
                            import traceback
                            with st.expander("Error Details"):
                                st.code(traceback.format_exc())
                        
                        finally:
                            st.session_state.extraction_in_progress = False
            
            if st.session_state.extraction_result:
                st.divider()
                result = st.session_state.extraction_result

                col_a, col_b = st.columns(2)
                with col_a:
                    completeness = result.get("completeness_score", 0)
                    st.metric("Completeness", f"{completeness}%")
                with col_b:
                    entities_count = len(result.get("layer1_finbert", {}).get("entities", []))
                    st.metric("Entities", entities_count)
                
                st.subheader("✨ Extracted Information")
                formatted_output = result.get("layer3_formatted", "No output")
                st.text_area("Entities:", value=formatted_output, height=400, disabled=True)
                
                st.download_button(
                    label="Download Extraction",
                    data=formatted_output,
                    file_name=f"extraction_{uploaded_file.name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain",
                    use_container_width=True
                )
        
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            import traceback
            with st.expander("Error Details"):
                st.code(traceback.format_exc())
    
    else:
        st.info("Upload a financial document (PDF, Word, Text, or Image) to start")


with col2:
    st.header("📊 Statistics")
    
    if st.session_state.extraction_result:
        result = st.session_state.extraction_result
        
        st.subheader("Quality Metrics")
        completeness = result.get("completeness_score", 0)
        
        if completeness >= 95:
            st.success(f"✅ Excellent: {completeness}%")
        elif completeness >= 80:
            st.info(f"ℹ️ Good: {completeness}%")
        else:
            st.warning(f"⚠️ Fair: {completeness}%")
        
        st.divider()
        
        st.subheader("Entity Breakdown")
        entities = result.get("layer1_finbert", {}).get("entities", [])
        
        if entities:
            entity_types = {}
            for entity in entities:
                etype = entity.get("type", "Unknown")
                entity_types[etype] = entity_types.get(etype, 0) + 1
            
            for etype, count in sorted(entity_types.items(), key=lambda x: x[1], reverse=True):
                st.metric(f"📌 {etype}", count)
        else:
            st.info("No entities detected")
        
        st.divider()
        
        st.subheader("Document Info")
        st.info(f"""
        **File:** {st.session_state.last_extracted_file}
        
        **Total Entities:** {len(entities)}
        
        **Model:** FinBERT + Gemini 2.5 Flash
        """)
    
    else:
        st.info("👈 Extract a document to see statistics")
        

st.divider()
